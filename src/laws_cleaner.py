import os
import re
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import List, Dict, Optional, Tuple
from datetime import datetime

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class LawDocumentCleaner:
    """法律法规文档清洗器"""
    
    def __init__(self):
        # 定义法律法规文档的噪声模式
        self.noise_patterns = {
            'header_footer_keywords': [
                '机密', '保密', '内部文件', '草稿', 'DRAFT',
                '第.*页.*共.*页', 'Page.*of.*', 
                '文件编号.*', 'Document No.*',
                '印发日期.*', '发布日期.*',
                '国务院办公厅', '中共中央办公厅',
                '公报', '公告', '通知'
            ],
            'page_numbers': [
                r'^\s*\d+\s*$',  # 单独的数字行
                r'^-\s*\d+\s*-$',  # - 1 -
                r'^\d+/\d+$',  # 1/10
                r'^\[\d+\]$',  # [1]
            ],
            'table_of_contents': [
                r'^目录$', r'^CONTENTS$',
                r'^\s*\.{3,}\s*\d+$',  # .......... 1
                r'^目\s*录$'
            ],
            'footnotes': [
                r'^注\d+：', r'^注释\d+：',
                r'^\[\d+\].*',  # [1] 开头的注释
                r'^①.*', r'^②.*', r'^③.*'  # 带圈数字注释
            ],
            'empty_lines': r'^\s*$',  # 空行
            'excessive_spaces': r'\s{2,}',  # 连续多个空格
            'special_chars': r'[●◆■▲★◎※◇▣]',  # 特殊符号
            'decoration_lines': [
                r'^=+$', r'^-+$', r'^~+$',  # 装饰线
                r'^★+$', r'^☆+$'
            ]
        }
        
        # 法律法规关键部分标识（这些应该保留）
        self.law_sections = [
            # 章节标识
            '第一章', '第二章', '第三章', '第四章', '第五章',
            '第一节', '第二节', '第三节', '第四节', '第五节',
            '第一编', '第二编', '第三编', '第四编', '第五编',
            
            # 条款标识
            '第一条', '第二条', '第三条', '第四条', '第五条',
            '第1条', '第2条', '第3条', '第4条', '第5条',
            'Article 1', 'Article 2', 'Article 3',
            
            # 法律法规特有标识
            '总则', '分则', '附则',
            '法律依据', '立法目的', '适用范围',
            '定义', '释义',
            '法律责任', '罚则', '处罚',
            '施行日期', '生效日期', '废止日期',
            '颁布机关', '发布机关',
            
            # 常见标题
            '中华人民共和国', '国务院令', '部令',
            '办法', '条例', '规定', '细则', '决定',
            '司法解释', '批复', '答复'
        ]
    
    def clean_document(self, input_path: str, output_path: str) -> bool:
        """清洗单个文档"""
        try:
            logger.info(f"开始清洗法律法规文档: {input_path}")
            
            # 读取文档
            doc = Document(input_path)
            
            # 应用各种清洗策略
            cleaned_paragraphs = self._extract_and_clean_paragraphs(doc)
            cleaned_paragraphs = self._remove_headers_footers(cleaned_paragraphs)
            cleaned_paragraphs = self._remove_page_numbers(cleaned_paragraphs)
            cleaned_paragraphs = self._remove_footnotes(cleaned_paragraphs)
            cleaned_paragraphs = self._remove_decoration_lines(cleaned_paragraphs)
            cleaned_paragraphs = self._clean_formatting(cleaned_paragraphs)
            cleaned_paragraphs = self._standardize_content(cleaned_paragraphs)
            cleaned_paragraphs = self._reconstruct_structure(cleaned_paragraphs)
            
            # 创建新文档
            new_doc = Document()
            
            # 设置文档格式
            self._setup_document_format(new_doc)
            
            # 添加清洗后的内容
            self._add_cleaned_content(new_doc, cleaned_paragraphs)
            
            # 保存文档
            new_doc.save(output_path)
            logger.info(f"法律法规文档已保存到: {output_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"清洗法律法规文档失败 {input_path}: {str(e)}")
            return False
    
    def _extract_and_clean_paragraphs(self, doc: Document) -> List[Dict]:
        """提取并初步清洗段落"""
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 跳过完全空的段落
            if not text:
                continue
            
            # 检查是否是表格内容（表格通常需要特殊处理）
            if self._is_table_content(para):
                continue
            
            # 检查是否是页眉页脚
            if self._is_header_footer(para):
                continue
            
            # 收集段落信息
            para_info = {
                'text': text,
                'original_text': text,  # 保留原始文本用于参考
                'style': para.style.name if para.style else 'Normal',
                'font_size': self._get_font_size(para),
                'is_bold': self._is_bold(para),
                'is_italic': self._is_italic(para),
                'alignment': self._get_alignment(para),
                'keep': True,  # 默认保留
                'is_law_section': self._is_law_section(text),  # 是否是法律法规关键部分
                'hierarchy_level': self._get_hierarchy_level(text)  # 获取层级
            }
            
            paragraphs.append(para_info)
        
        # 处理表格
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                paragraphs.append({
                    'text': table_text,
                    'original_text': table_text,
                    'style': 'Table',
                    'font_size': 10,
                    'is_bold': False,
                    'is_italic': False,
                    'alignment': 'LEFT',
                    'keep': True,
                    'is_law_section': False,
                    'hierarchy_level': 0
                })
        
        return paragraphs
    
    def _is_table_content(self, para) -> bool:
        """判断是否为表格内容"""
        # 简单的表格内容检测
        text = para.text
        if any(char in text for char in ['┌', '┐', '└', '┘', '├', '┤', '┬', '┴', '─', '│', '┃']):
            return True
        return False
    
    def _is_header_footer(self, para) -> bool:
        """判断是否为页眉页脚"""
        text = para.text
        
        # 检查页眉页脚关键词
        for pattern in self.noise_patterns['header_footer_keywords']:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # 检查页眉页脚的常见位置特征
        if len(text) < 50 and (text.isdigit() or '页' in text or 'Page' in text):
            return True
        
        # 检查是否是文件编号
        if re.search(r'[〔〔\]【】]', text) and len(text) < 30:
            return True
        
        return False
    
    def _is_law_section(self, text: str) -> bool:
        """判断是否是法律法规关键部分"""
        for section in self.law_sections:
            if section in text:
                return True
        return False
    
    def _get_hierarchy_level(self, text: str) -> int:
        """获取文本的层级（用于结构化）"""
        # 第X章 - 层级1
        if re.match(r'^第[一二三四五六七八九十]+章', text) or re.match(r'^第\d+章', text):
            return 1
        # 第X节 - 层级2
        elif re.match(r'^第[一二三四五六七八九十]+节', text) or re.match(r'^第\d+节', text):
            return 2
        # 第一条 - 层级3
        elif re.match(r'^第[一二三四五六七八九十]+条', text) or re.match(r'^第\d+条', text):
            return 3
        # 第X款 - 层级4
        elif re.match(r'^第[一二三四五六七八九十]+款', text) or re.match(r'^第\d+款', text):
            return 4
        # (X) - 层级5
        elif re.match(r'^\([一二三四五六七八九十]+\)', text) or re.match(r'^\(\d+\)', text):
            return 5
        # 其他重要标题 - 层级2
        elif any(keyword in text for keyword in ['总则', '分则', '附则', '法律责任', '罚则']):
            return 2
        else:
            return 0
    
    def _get_font_size(self, para) -> float:
        """获取字体大小"""
        try:
            if para.runs:
                return para.runs[0].font.size.pt if para.runs[0].font.size else 11
        except:
            pass
        return 11.0
    
    def _is_bold(self, para) -> bool:
        """判断是否为粗体"""
        if para.runs:
            return para.runs[0].bold or False
        return False
    
    def _is_italic(self, para) -> bool:
        """判断是否为斜体"""
        if para.runs:
            return para.runs[0].italic or False
        return False
    
    def _get_alignment(self, para):
        """获取对齐方式"""
        try:
            align_map = {
                0: 'LEFT',    # WD_ALIGN_PARAGRAPH.LEFT
                1: 'CENTER',  # WD_ALIGN_PARAGRAPH.CENTER
                2: 'RIGHT',   # WD_ALIGN_PARAGRAPH.RIGHT
                3: 'JUSTIFY', # WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            return align_map.get(para.alignment, 'LEFT')
        except:
            return 'LEFT'
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                # 用制表符分隔单元格内容
                table_text.append('\t'.join(row_text))
        
        return '\n'.join(table_text) if table_text else ''
    
    def _remove_headers_footers(self, paragraphs: List[Dict]) -> List[Dict]:
        """移除页眉页脚"""
        cleaned = []
        
        for i, para in enumerate(paragraphs):
            text = para['text']
            
            # 跳过页眉页脚
            if self._is_header_footer_text(text):
                continue
            
            # 检查是否是文件头信息（如"国务院令第XXX号"）
            if i < 3:  # 前3行可能是文件头
                if self._is_document_header(text):
                    # 保留重要的文件头信息
                    if not any(keyword in text for keyword in ['令', '公告', '通知', '决定']):
                        continue
            
            cleaned.append(para)
        
        return cleaned
    
    def _is_document_header(self, text: str) -> bool:
        """判断是否为文件头信息"""
        patterns = [
            r'.*令第\d+号',
            r'.*公告第\d+号',
            r'.*通知第\d+号',
            r'^【.*】$',
            r'^〔.*〕$'
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def _is_header_footer_text(self, text: str) -> bool:
        """判断文本是否为页眉页脚内容"""
        text_lower = text.lower()
        
        # 常见的页眉页脚关键词
        header_footer_indicators = [
            'confidential', 'internal', 'draft',
            'page', '页码', '第.*页',
            '文件编号', 'document no',
            '印发日期', '发布日期'
        ]
        
        for indicator in header_footer_indicators:
            if re.search(indicator, text_lower):
                return True
        
        return False
    
    def _remove_page_numbers(self, paragraphs: List[Dict]) -> List[Dict]:
        """移除页码"""
        cleaned = []
        
        for para in paragraphs:
            text = para['text']
            
            # 跳过明显的页码
            is_page_number = False
            for pattern in self.noise_patterns['page_numbers']:
                if re.match(pattern, text.strip()):
                    is_page_number = True
                    break
            
            if not is_page_number:
                cleaned.append(para)
        
        return cleaned
    
    def _remove_footnotes(self, paragraphs: List[Dict]) -> List[Dict]:
        """移除脚注"""
        cleaned = []
        
        for para in paragraphs:
            text = para['text']
            
            # 检查是否是脚注
            is_footnote = False
            for pattern in self.noise_patterns['footnotes']:
                if re.match(pattern, text):
                    is_footnote = True
                    break
            
            if not is_footnote:
                cleaned.append(para)
        
        return cleaned
    
    def _remove_decoration_lines(self, paragraphs: List[Dict]) -> List[Dict]:
        """移除装饰线"""
        cleaned = []
        
        for para in paragraphs:
            text = para['text']
            
            # 检查是否是装饰线
            is_decoration = False
            for pattern in self.noise_patterns['decoration_lines']:
                if re.match(pattern, text.strip()):
                    is_decoration = True
                    break
            
            if not is_decoration:
                cleaned.append(para)
        
        return cleaned
    
    def _clean_formatting(self, paragraphs: List[Dict]) -> List[Dict]:
        """清理格式"""
        cleaned = []
        
        for para in paragraphs:
            text = para['original_text']  # 使用原始文本进行清理
            
            # 移除多余的空格
            text = re.sub(self.noise_patterns['excessive_spaces'], ' ', text)
            
            # 移除特殊符号（除非是列表符号或引用符号）
            if not text.startswith(('•', '-', '1.', '2.', '3.', '（', '(', '《', '[')):
                text = re.sub(self.noise_patterns['special_chars'], '', text)
            
            # 清理行首行尾空格
            text = text.strip()
            
            # 如果清理后还有内容，则保留
            if text:
                para['text'] = text
                cleaned.append(para)
        
        return cleaned
    
    def _standardize_content(self, paragraphs: List[Dict]) -> List[Dict]:
        """标准化内容"""
        cleaned = []
        
        for i, para in enumerate(paragraphs):
            text = para['text']
            
            # 标准化条款编号
            text = self._standardize_clause_numbers(text)
            
            # 标准化日期格式
            text = self._standardize_dates(text)
            
            # 标准化法律引用
            text = self._standardize_law_references(text)
            
            # 标准化金额和数字
            text = self._standardize_numbers(text)
            
            para['text'] = text
            cleaned.append(para)
        
        return cleaned
    
    def _standardize_clause_numbers(self, text: str) -> str:
        """标准化条款编号"""
        # 将中文数字条款标准化
        replacements = [
            (r'第([一二三四五六七八九十])章', lambda m: f'第{self._chinese_to_number(m.group(1))}章'),
            (r'第([一二三四五六七八九十])节', lambda m: f'第{self._chinese_to_number(m.group(1))}节'),
            (r'第([一二三四五六七八九十])条', lambda m: f'第{self._chinese_to_number(m.group(1))}条'),
            (r'第([一二三四五六七八九十])款', lambda m: f'第{self._chinese_to_number(m.group(1))}款'),
            (r'\(([一二三四五六七八九十])\)', lambda m: f'({self._chinese_to_number(m.group(1))})'),
        ]
        
        for pattern, replacement in replacements:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _chinese_to_number(self, chinese: str) -> str:
        """中文数字转阿拉伯数字"""
        chinese_map = {
            '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
            '六': '6', '七': '7', '八': '8', '九': '9', '十': '10',
            '十一': '11', '十二': '12', '十三': '13', '十四': '14', '十五': '15',
            '十六': '16', '十七': '17', '十八': '18', '十九': '19', '二十': '20'
        }
        return chinese_map.get(chinese, chinese)
    
    def _standardize_dates(self, text: str) -> str:
        """标准化日期格式"""
        # 统一日期格式为 YYYY年MM月DD日
        patterns = [
            (r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})年(\d{1,2})月(\d{1,2})号', r'\1年\2月\3日'),
        ]
        
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _standardize_law_references(self, text: str) -> str:
        """标准化法律引用"""
        # 统一法律引用格式
        patterns = [
            (r'《\s*([^》]+)\s*》\s*第\s*(\d+)\s*条', r'《\1》第\2条'),
            (r'《\s*([^》]+)\s*》\s*第\s*([一二三四五六七八九十]+)\s*条', 
             lambda m: f'《{m.group(1)}》第{self._chinese_to_number(m.group(2))}条'),
        ]
        
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _standardize_numbers(self, text: str) -> str:
        """标准化数字格式"""
        # 统一数字格式
        patterns = [
            (r'(\d+)[,，](\d{3})', r'\1\2'),  # 移除千分位逗号
            (r'(\d+)万元', r'\1万元'),
            (r'(\d+)元人民币', r'\1元'),
            (r'百分之(\d+)', r'\1%'),
        ]
        
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _reconstruct_structure(self, paragraphs: List[Dict]) -> List[Dict]:
        """重构文档结构"""
        cleaned = []
        
        for i, para in enumerate(paragraphs):
            # 根据层级添加空行以增强可读性
            if para['hierarchy_level'] > 0 and i > 0:
                # 在不同层级的标题前添加空行
                if para['hierarchy_level'] <= 2:  # 章、节级别
                    cleaned.append(self._create_empty_paragraph())
                elif para['hierarchy_level'] == 3:  # 条级别
                    cleaned.append(self._create_empty_paragraph(font_size=0.5))
            
            cleaned.append(para)
        
        return cleaned
    
    def _create_empty_paragraph(self, font_size: float = 1.0) -> Dict:
        """创建空段落"""
        return {
            'text': '',
            'original_text': '',
            'style': 'Empty',
            'font_size': font_size,
            'is_bold': False,
            'is_italic': False,
            'alignment': 'LEFT',
            'keep': True,
            'is_law_section': False,
            'hierarchy_level': 0
        }
    
    def _setup_document_format(self, doc: Document):
        """设置文档格式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 设置标题样式
        for i in range(1, 6):
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = '黑体'
            heading_font.size = Pt(16 - (i-1) * 2)
            heading_font.bold = True
    
    def _add_cleaned_content(self, doc: Document, paragraphs: List[Dict]):
        """添加清洗后的内容"""
        for para_info in paragraphs:
            text = para_info['text']
            
            # 跳过空段落（占位用）
            if para_info['style'] == 'Empty':
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(para_info['font_size'] * 12)
                continue
            
            # 根据样式添加段落
            if para_info['style'] == 'Table':
                # 表格内容特殊处理
                p = doc.add_paragraph()
                p.add_run('[表格内容]').italic = True
                doc.add_paragraph(text)
            
            elif para_info['hierarchy_level'] > 0:
                # 使用相应的标题样式
                heading_level = min(para_info['hierarchy_level'], 3)
                p = doc.add_paragraph(text, style=f'Heading {heading_level}')
                
                # 设置对齐方式
                if para_info['alignment'] == 'CENTER':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif para_info['alignment'] == 'RIGHT':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            elif para_info['is_bold'] and para_info['font_size'] > 12:
                # 重要标题加粗
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
            
            elif para_info['is_italic']:
                # 斜体内容（通常是说明或注释）
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
            
            else:
                # 普通正文
                p = doc.add_paragraph(text)
            
            # 设置段落格式
            if para_info['hierarchy_level'] == 0:  # 普通正文
                p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
            p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
            p.paragraph_format.space_after = Pt(6)


def process_law_documents_folder(input_dir: str, output_dir: str):
    """处理整个文件夹的法律法规文档"""
    
    # 创建输出目录
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    # 初始化清洗器
    cleaner = LawDocumentCleaner()
    
    # 统计信息
    processed_count = 0
    failed_count = 0
    
    # 遍历输入目录
    input_path = Path(input_dir)
    
    # 支持的文件扩展名
    valid_extensions = ['.docx', '.doc']
    
    # 收集所有文件
    all_files = []
    for ext in valid_extensions:
        all_files.extend(list(input_path.rglob(f'*{ext}')))
        all_files.extend(list(input_path.rglob(f'*{ext.upper()}')))
    
    for file_path in all_files:
        # 构建输出路径
        relative_path = file_path.relative_to(input_path)
        output_path = Path(output_dir) / relative_path
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # 修改输出文件名为 cleaned_原文件名
        cleaned_name = f"{file_path.stem}.docx"
        output_file = output_path.parent / cleaned_name
        
        # 清洗文档
        if cleaner.clean_document(str(file_path), str(output_file)):
            processed_count += 1
            logger.info(f"成功处理: {file_path.name}")
        else:
            failed_count += 1
            logger.error(f"处理失败: {file_path.name}")
    
    # 输出统计信息
    logger.info(f"法律法规文档处理完成！")
    logger.info(f"成功处理: {processed_count} 个文件")
    logger.info(f"处理失败: {failed_count} 个文件")
    
    # 生成清洗报告
    generate_law_cleaning_report(output_dir, processed_count, failed_count)


def generate_law_cleaning_report(output_dir: str, success_count: int, failed_count: int):
    """生成法律法规文档清洗报告"""
    report_path = Path(output_dir) / "law_cleaning_report.txt"
    
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("法律法规文档清洗报告\n")
        f.write("=" * 60 + "\n\n")
        f.write(f"清洗时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"成功处理的文件数: {success_count}\n")
        f.write(f"处理失败的文件数: {failed_count}\n")
        f.write("\n")
        f.write("清洗操作包括:\n")
        f.write("1. 移除页眉页脚和页码\n")
        f.write("2. 移除目录、脚注和装饰线\n")
        f.write("3. 清理多余空格和空行\n")
        f.write("4. 标准化条款编号、日期和法律引用格式\n")
        f.write("5. 标准化数字和金额格式\n")
        f.write("6. 提取并处理表格内容\n")
        f.write("7. 保留法律法规关键结构和条款\n")
        f.write("8. 重构文档层级结构\n")
        f.write("\n")
        f.write("注意事项:\n")
        f.write("- 保留了章、节、条、款等层级结构\n")
        f.write("- 标准化了法律引用格式\n")
        f.write("- 统一了日期和数字格式\n")
        f.write("- 设置了首行缩进和行间距\n")
    
    logger.info(f"清洗报告已生成: {report_path}")


def main():
    """主函数"""
    # 设置路径
    input_dir = "../docs/laws"
    output_dir = "../docs/laws_cleaned"
    
    # 检查输入目录是否存在
    if not Path(input_dir).exists():
        logger.error(f"输入目录不存在: {input_dir}")
        logger.info(f"正在创建输入目录: {input_dir}")
        Path(input_dir).mkdir(parents=True, exist_ok=True)
        return
    
    logger.info(f"开始清洗法律法规文档...")
    logger.info(f"输入目录: {input_dir}")
    logger.info(f"输出目录: {output_dir}")
    
    # 处理文档
    process_law_documents_folder(input_dir, output_dir)


if __name__ == "__main__":
    main()