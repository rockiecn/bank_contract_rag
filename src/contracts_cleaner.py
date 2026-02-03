import os
import re
import shutil
from pathlib import Path
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import List, Dict, Optional, Tuple

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ContractCleaner:
    """合同文档清洗器"""
    
    def __init__(self):
        # 定义常见的合同噪声模式
        self.noise_patterns = {
            'header_footer_keywords': [
                '机密', '保密', '内部文件', '草稿', 'DRAFT',
                '第.*页.*共.*页', 'Page.*of.*', 
                '合同编号.*', 'Contract No.*',
                '日期.*', 'Date.*'
            ],
            'page_numbers': [
                r'^\s*\d+\s*$',  # 单独的数字行
                r'^-\s*\d+\s*-$',  # - 1 -
                r'^\d+/\d+$',  # 1/10
            ],
            'table_of_contents': [
                r'^目录$', r'^CONTENTS$',
                r'^\s*\.{3,}\s*\d+$',  # .......... 1
            ],
            'empty_lines': r'^\s*$',  # 空行
            'excessive_spaces': r'\s{2,}',  # 连续多个空格
            'special_chars': r'[●◆■▲★◎※◇]',  # 特殊符号
        }
        
        # 合同关键部分标识（这些应该保留）
        self.contract_sections = [
            '甲方', '乙方', '丙方', '丁方',
            '第一条', '第二条', '第三条', '第四条', '第五条',
            '第1条', '第2条', '第3条', '第4条', '第5条',
            'Article 1', 'Article 2', 'Article 3',
            '第一章', '第二章', '第三章',
            '合同双方', '合同标的', '合同价款', '付款方式',
            '违约责任', '争议解决', '保密条款',
            '签字页', '签署', '盖章'
        ]
    
    def clean_document(self, input_path: str, output_path: str) -> bool:
        """清洗单个文档"""
        try:
            logger.info(f"开始清洗文档: {input_path}")
            
            # 读取文档
            doc = Document(input_path)
            
            # 应用各种清洗策略
            cleaned_paragraphs = self._extract_and_clean_paragraphs(doc)
            cleaned_paragraphs = self._remove_headers_footers(cleaned_paragraphs)
            cleaned_paragraphs = self._remove_page_numbers(cleaned_paragraphs)
            cleaned_paragraphs = self._clean_formatting(cleaned_paragraphs)
            cleaned_paragraphs = self._standardize_content(cleaned_paragraphs)
            
            # 创建新文档
            new_doc = Document()
            
            # 设置文档格式
            self._setup_document_format(new_doc)
            
            # 添加清洗后的内容
            self._add_cleaned_content(new_doc, cleaned_paragraphs)
            
            # 保存文档
            new_doc.save(output_path)
            logger.info(f"文档已保存到: {output_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"清洗文档失败 {input_path}: {str(e)}")
            return False
    
    def _extract_and_clean_paragraphs(self, doc: Document) -> List[Dict]:
        """提取并初步清洗段落"""
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 跳过完全空的段落
            if not text:
                continue
            
            # 检查是否是表格内容（表格通常需要特殊处理）
            if self._is_table_content(para):
                continue
            
            # 检查是否是页眉页脚
            if self._is_header_footer(para):
                continue
            
            # 收集段落信息
            para_info = {
                'text': text,
                'style': para.style.name if para.style else 'Normal',
                'font_size': self._get_font_size(para),
                'is_bold': self._is_bold(para),
                'alignment': self._get_alignment(para),
                'keep': True  # 默认保留
            }
            
            paragraphs.append(para_info)
        
        # 处理表格
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                paragraphs.append({
                    'text': table_text,
                    'style': 'Table',
                    'font_size': 10,
                    'is_bold': False,
                    'alignment': 'LEFT',
                    'keep': True
                })
        
        return paragraphs
    
    def _is_table_content(self, para) -> bool:
        """判断是否为表格内容"""
        # 简单的表格内容检测
        text = para.text
        if any(char in text for char in ['┌', '┐', '└', '┘', '├', '┤', '┬', '┴', '─', '│']):
            return True
        return False
    
    def _is_header_footer(self, para) -> bool:
        """判断是否为页眉页脚"""
        text = para.text
        
        # 检查页眉页脚关键词
        for pattern in self.noise_patterns['header_footer_keywords']:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # 检查页眉页脚的常见位置特征
        if len(text) < 50 and (text.isdigit() or '页' in text or 'Page' in text):
            return True
        
        return False
    
    def _get_font_size(self, para) -> float:
        """获取字体大小"""
        try:
            if para.runs:
                return para.runs[0].font.size.pt if para.runs[0].font.size else 11
        except:
            pass
        return 11.0
    
    def _is_bold(self, para) -> bool:
        """判断是否为粗体"""
        if para.runs:
            return para.runs[0].bold or False
        return False
    
    def _get_alignment(self, para):
        """获取对齐方式"""
        try:
            align_map = {
                0: 'LEFT',    # WD_ALIGN_PARAGRAPH.LEFT
                1: 'CENTER',  # WD_ALIGN_PARAGRAPH.CENTER
                2: 'RIGHT',   # WD_ALIGN_PARAGRAPH.RIGHT
                3: 'JUSTIFY', # WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            return align_map.get(para.alignment, 'LEFT')
        except:
            return 'LEFT'
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                # 用制表符分隔单元格内容
                table_text.append('\t'.join(row_text))
        
        return '\n'.join(table_text) if table_text else ''
    
    def _remove_headers_footers(self, paragraphs: List[Dict]) -> List[Dict]:
        """移除页眉页脚"""
        cleaned = []
        
        for i, para in enumerate(paragraphs):
            text = para['text']
            
            # 跳过页眉页脚
            if self._is_header_footer_text(text):
                continue
            
            # 检查上下文，避免误删合同标题
            if i > 0 and i < len(paragraphs) - 1:
                prev_text = paragraphs[i-1]['text']
                next_text = paragraphs[i+1]['text']
                
                # 如果是单独的数字行，且上下文不是条款，则很可能是页码
                if re.match(r'^\s*\d+\s*$', text):
                    if not any(section in prev_text or section in next_text 
                              for section in self.contract_sections):
                        continue
            
            cleaned.append(para)
        
        return cleaned
    
    def _is_header_footer_text(self, text: str) -> bool:
        """判断文本是否为页眉页脚内容"""
        text_lower = text.lower()
        
        # 常见的页眉页脚关键词
        header_footer_indicators = [
            'confidential', 'internal', 'draft',
            'page', '页码', '第.*页',
            '合同编号', 'contract no',
            '日期', 'date'
        ]
        
        for indicator in header_footer_indicators:
            if re.search(indicator, text_lower):
                return True
        
        return False
    
    def _remove_page_numbers(self, paragraphs: List[Dict]) -> List[Dict]:
        """移除页码"""
        cleaned = []
        
        for para in paragraphs:
            text = para['text']
            
            # 跳过明显的页码
            is_page_number = False
            for pattern in self.noise_patterns['page_numbers']:
                if re.match(pattern, text.strip()):
                    is_page_number = True
                    break
            
            if not is_page_number:
                cleaned.append(para)
        
        return cleaned
    
    def _clean_formatting(self, paragraphs: List[Dict]) -> List[Dict]:
        """清理格式"""
        cleaned = []
        
        for para in paragraphs:
            text = para['text']
            
            # 移除多余的空格
            text = re.sub(self.noise_patterns['excessive_spaces'], ' ', text)
            
            # 移除特殊符号（除非是列表符号）
            if not text.startswith(('•', '-', '1.', '2.', '3.', '（', '(')):
                text = re.sub(self.noise_patterns['special_chars'], '', text)
            
            # 清理行首行尾空格
            text = text.strip()
            
            # 如果清理后还有内容，则保留
            if text:
                para['text'] = text
                cleaned.append(para)
        
        return cleaned
    
    def _standardize_content(self, paragraphs: List[Dict]) -> List[Dict]:
        """标准化内容"""
        cleaned = []
        
        for i, para in enumerate(paragraphs):
            text = para['text']
            
            # 标准化条款编号
            text = self._standardize_clause_numbers(text)
            
            # 标准化日期格式
            text = self._standardize_dates(text)
            
            # 标准化金额格式
            text = self._standardize_amounts(text)
            
            para['text'] = text
            cleaned.append(para)
        
        return cleaned
    
    def _standardize_clause_numbers(self, text: str) -> str:
        """标准化条款编号"""
        # 将中文数字条款标准化
        replacements = {
            r'第(\d+)条': r'第\1条',
            r'第([一二三四五六七八九十])条': lambda m: f'第{self._chinese_to_number(m.group(1))}条',
            r'Article\s*(\d+)': r'Article \1',
        }
        
        for pattern, replacement in replacements.items():
            if callable(replacement):
                text = re.sub(pattern, replacement, text)
            else:
                text = re.sub(pattern, replacement, text)
        
        return text
    
    def _chinese_to_number(self, chinese: str) -> str:
        """中文数字转阿拉伯数字"""
        chinese_map = {
            '一': '1', '二': '2', '三': '3', '四': '4', '五': '5',
            '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'
        }
        return chinese_map.get(chinese, chinese)
    
    def _standardize_dates(self, text: str) -> str:
        """标准化日期格式"""
        # 统一日期格式为 YYYY年MM月DD日
        patterns = [
            (r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})', r'\1年\2月\3日'),
            (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1年\2月\3日'),
        ]
        
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _standardize_amounts(self, text: str) -> str:
        """标准化金额格式"""
        # 统一金额格式
        patterns = [
            (r'(\d+)[,，](\d{3})', r'\1\2'),  # 移除千分位逗号
            (r'人民币\s*(\d+(?:\.\d+)?)\s*元', r'人民币\1元'),
            (r'RMB\s*(\d+(?:\.\d+)?)\s*yuan', r'人民币\1元'),
        ]
        
        for pattern, replacement in patterns:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _setup_document_format(self, doc: Document):
        """设置文档格式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
    
    def _add_cleaned_content(self, doc: Document, paragraphs: List[Dict]):
        """添加清洗后的内容"""
        for para_info in paragraphs:
            text = para_info['text']
            
            # 根据样式添加段落
            if para_info['style'] == 'Table':
                # 表格内容特殊处理
                p = doc.add_paragraph()
                p.add_run('[表格内容]').italic = True
                doc.add_paragraph(text)
            elif para_info['is_bold'] or any(keyword in text for keyword in ['甲方', '乙方', '第一条']):
                # 重要内容加粗
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            else:
                doc.add_paragraph(text)
            
            # 添加段落间距
            if len(doc.paragraphs) > 1:
                doc.paragraphs[-2].paragraph_format.space_after = Pt(6)


def process_contracts_folder(input_dir: str, output_dir: str):
    """处理整个文件夹的合同文档"""
    
    # 创建输出目录
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    # 初始化清洗器
    cleaner = ContractCleaner()
    
    # 统计信息
    processed_count = 0
    failed_count = 0
    
    # 遍历输入目录
    input_path = Path(input_dir)
    
    # 支持的文件扩展名
    valid_extensions = ['.docx', '.doc']
    
    for file_path in input_path.rglob('*'):
        if file_path.suffix.lower() in valid_extensions:
            # 构建输出路径
            relative_path = file_path.relative_to(input_path)
            output_path = Path(output_dir) / relative_path
            
            # 确保输出目录存在
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 修改输出文件名为 cleaned_原文件名
            cleaned_name = f"{file_path.stem}.docx"
            output_file = output_path.parent / cleaned_name
            
            # 清洗文档
            if cleaner.clean_document(str(file_path), str(output_file)):
                processed_count += 1
                logger.info(f"成功处理: {file_path.name}")
            else:
                failed_count += 1
                logger.error(f"处理失败: {file_path.name}")
    
    # 输出统计信息
    logger.info(f"处理完成！")
    logger.info(f"成功处理: {processed_count} 个文件")
    logger.info(f"处理失败: {failed_count} 个文件")
    
    # 生成清洗报告
    generate_cleaning_report(output_dir, processed_count, failed_count)


def generate_cleaning_report(output_dir: str, success_count: int, failed_count: int):
    """生成清洗报告"""
    report_path = Path(output_dir) / "cleaning_report.txt"
    
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("合同文档清洗报告\n")
        f.write("=" * 50 + "\n\n")
        f.write(f"清洗时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"成功处理的文件数: {success_count}\n")
        f.write(f"处理失败的文件数: {failed_count}\n")
        f.write("\n")
        f.write("清洗操作包括:\n")
        f.write("1. 移除页眉页脚和页码\n")
        f.write("2. 移除目录和特殊符号\n")
        f.write("3. 清理多余空格和空行\n")
        f.write("4. 标准化条款编号、日期和金额格式\n")
        f.write("5. 提取并处理表格内容\n")
        f.write("6. 保留合同关键条款和结构\n")
    
    logger.info(f"清洗报告已生成: {report_path}")


def main():
    """主函数"""
    # 设置路径
    input_dir = "../docs/contracts"
    output_dir = "../docs/contracts_cleaned"
    
    # 检查输入目录是否存在
    if not Path(input_dir).exists():
        logger.error(f"输入目录不存在: {input_dir}")
        return
    
    logger.info(f"开始清洗合同文档...")
    logger.info(f"输入目录: {input_dir}")
    logger.info(f"输出目录: {output_dir}")
    
    # 处理文档
    process_contracts_folder(input_dir, output_dir)


if __name__ == "__main__":
    from datetime import datetime
    main()